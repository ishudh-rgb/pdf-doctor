#!/usr/bin/env python3
from __future__ import annotations

import re
import sys
from pathlib import Path

import fitz
from docx import Document
from docx.oxml.ns import qn


def pdf_watermarks(pdf_path: str) -> None:
    doc = fitz.open(pdf_path)
    for page_idx in range(min(3, len(doc))):
        page = doc[page_idx]
        print(f"\n=== PDF page {page_idx + 1} ({page.rect.width}x{page.rect.height}) ===")
        for block in page.get_text("dict")["blocks"]:
            if block.get("type") != 0:
                continue
            for line in block["lines"]:
                text = "".join(span["text"] for span in line["spans"]).strip()
                if "CONFIDENTIAL" in text:
                    sizes = [round(span["size"], 1) for span in line["spans"]]
                    print(f"  CONF bbox={tuple(round(x, 1) for x in line['bbox'])} sizes={sizes}")


def docx_page1(docx_path: str) -> None:
    doc = Document(docx_path)
    print(f"\n=== DOCX first 30 paras ===")
    for i, para in enumerate(doc.paragraphs[:30]):
        text = para.text.strip()
        p_pr = para._p.find(qn("w:pPr"))
        num = None
        if p_pr is not None:
            num_pr = p_pr.find(qn("w:numPr"))
            if num_pr is not None:
                nid = num_pr.find(qn("w:numId"))
                num = nid.get(qn("w:val")) if nid is not None else "?"
        wm = "WM" if "CONFIDENTIAL" in para._p.xml else ""
        if not text:
            if num or wm:
                print(f"  {i:3d} EMPTY num={num} {wm}")
            continue
        print(f"  {i:3d} n={num or '-':2s} {wm:2s} {text[:75]}")

    # find solution split
    print("\n=== Solution section paras ===")
    for i, para in enumerate(doc.paragraphs):
        t = para.text.strip()
        if "Solution" in t or "Contract manufacturing" in t or "accessible-premium" in t:
            num = para._p.find(qn("w:pPr"))
            has_num = num is not None and num.find(qn("w:numPr")) is not None
            print(f"  {i:3d} bullet={has_num} {t[:90]}")


def main() -> None:
    pdf_path = sys.argv[1]
    docx_path = sys.argv[2]
    pdf_watermarks(pdf_path)
    docx_page1(docx_path)


if __name__ == "__main__":
    main()
