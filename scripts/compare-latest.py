#!/usr/bin/env python3
from __future__ import annotations

import os
import sys
import zipfile
from docx import Document
from docx.oxml.ns import qn


def metrics(path: str) -> dict:
    with zipfile.ZipFile(path) as z:
        xml = z.read("word/document.xml").decode("utf-8", errors="ignore")
        media = [n for n in z.namelist() if n.startswith("word/media/")]
    doc = Document(path)
    conf = sum(1 for p in doc.paragraphs if "CONFIDENTIAL" in p.text)
    return {
        "file": os.path.basename(path),
        "size": os.path.getsize(path),
        "tables": xml.count("<w:tbl>"),
        "drawings": xml.count("<w:drawing"),
        "media": len(media),
        "numPr": xml.count("<w:numPr"),
        "conf_paras": conf,
        "conf_xml": xml.count("CONFIDENTIAL"),
    }


def problem_section(path: str) -> None:
    doc = Document(path)
    active = False
    print(f"\n--- Problem/Solution ({os.path.basename(path)}) ---")
    for p in doc.paragraphs:
        t = p.text.strip()
        if "Problem Statement" in t:
            active = True
        if not active:
            continue
        if t.startswith("Products"):
            break
        ppr = p._p.find(qn("w:pPr"))
        num = ppr is not None and ppr.find(qn("w:numPr")) is not None
        has_draw = bool(list(p._p.iter(qn("w:drawing"))))
        mark = "B" if num else " "
        extra = " [draw]" if has_draw else ""
        print(f"{mark}{extra} | {t[:85]}")


if __name__ == "__main__":
    for p in sys.argv[1:]:
        print(metrics(p))
        problem_section(p)
