#!/usr/bin/env python3
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn

d = Document(r"C:\Users\NiTiN Dhiman\Downloads\merged (3) (17).docx")
lines = []
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if any(
        k in t
        for k in (
            "Old working",
            "Page 2",
            "Page 3",
            "Page 5",
            "Marketing Model",
            "Financial Modeling",
            "Co-founder KRA",
        )
    ):
        lines.append(f"p{i} {t[:80]}")
lines.append(f"tables={len(d.tables)}")
for i, t in enumerate(d.tables):
    cols = len(t.columns)
    tx = "".join(c.text for r in t.rows for c in r.cells)
    if cols >= 4 or "Sl" in tx:
        grid = t._tbl.find(qn("w:tblGrid"))
        ws = [c.get(qn("w:w")) for c in grid.findall(qn("w:gridCol"))] if grid is not None else []
        tot = sum(int(w) for w in ws) if ws else 0
        lines.append(f"tbl{i} cols={cols} gridSum={tot} {tx[:50].replace(chr(10), ' ')}")
Path(r"c:\Users\NiTiN Dhiman\Documents\Cursor\pdf-doctor\scripts\ledger-scan.txt").write_text(
    "\n".join(lines), encoding="utf-8"
)
