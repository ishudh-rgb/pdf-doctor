"""Transform raw pdf2docx output to match Smallpdf result.docx layout."""
from __future__ import annotations

import re
import zipfile
from copy import deepcopy
from pathlib import Path

import fitz
from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

HEADING_RE = re.compile(r"^(\d+\.\s+)(.+)$")
CONTINUATION_START = re.compile(
    r"^[a-z]|^product choices|^sensitivity|^required\.|^hai\.|^control me|^rakho\."
)

TEMPLATE_DIR = Path(__file__).resolve().parent / "templates" / "smallpdf-ref"
WATERMARK_RUN_PATH = (
    Path(__file__).resolve().parent / "templates" / "confidential-watermark-vml.xml"
)
WATERMARK_DRAWING_PATH = (
    Path(__file__).resolve().parent / "templates" / "confidential-watermark-drawing.xml"
)
LEDGER_MARKERS = ("Sl. No.", "Debit(-)", "Credit(+)")
DECIMAL_NUM_ID = "1"  # 1. 2. 3. body lists (matches smallpdf numId 1)
BULLET_NUM_ID = "2"  # • bullets in SWOT table cells
MAX_WATERMARKS = 8
SWOT_HEADING = "6. SWOT Analysis"
SWOT_CELL_FILL = "FAFAFA"

WATERMARK_ANCHORS = (
    "Contract manufacturing use karo",
    "Formats",
    "50ml hero bottle: ₹1,799–₹2,490",
    "Business Model",
    "Sales Model",
    "Bootstrap launch capital:",
    "Recommended role: Co-founder",
    "Conversion & Retention — 15%",
)

# PDF list geometry — positions are from page edge; Word pgMar left is already 51pt (1020 twips)
PDF_PAGE_MARGIN_PT = 51.0
BULLET_TEXT_LEFT_PT = 65.8
SECTION_LEFT_PT = 51.0
BULLET_OFFSET_PT = BULLET_TEXT_LEFT_PT - PDF_PAGE_MARGIN_PT  # 14.8pt past margin
# Hanging: bullet at margin, wrapped text aligns at BULLET_OFFSET
BULLET_IND_LEFT = str(int(round(BULLET_OFFSET_PT * 20)))
BULLET_IND_HANGING = BULLET_IND_LEFT
SECTION_IND_LEFT = "0"
BODY_FONT_HALFPTS = "22"   # 11pt
HEADING_FONT_HALFPTS = "44"  # 22pt
# PDF CONFIDENTIAL bbox (pt) — same on pages 1–8 of merged (3).pdf
WATERMARK_BBOX = (110.37, 378.89, 440.98, 637.98)
PAGE_WIDTH_PT = 595.276
# Nudge left + widen so rotated "L" is not clipped at the right margin
WATERMARK_MARGIN_LEFT_PT = WATERMARK_BBOX[0] - 32
WATERMARK_WIDTH_PT = (WATERMARK_BBOX[2] - WATERMARK_BBOX[0]) + 48
WATERMARK_TOP_PT = WATERMARK_BBOX[1]
# A4 twips: portrait content width = page - left margin - right margin
PORTRAIT_CONTENT_TWIPS = 11906 - 1020 - 1316
LANDSCAPE_CONTENT_TWIPS = 16838 - 1020 - 1316
LEDGER_START_PREFIX = "Old working Sheet"


def _norm_key(text: str) -> str:
    return re.sub(r"\s+", " ", text.strip().lower())[:96]


def build_pdf_list_map(pdf_path: str) -> tuple[set[str], set[str]]:
    """Map PDF x-position list structure: main headings (decimal) vs sub-items (dot)."""
    decimal_titles: set[str] = set()
    bullet_keys: set[str] = set()
    doc = fitz.open(pdf_path)
    try:
        for page in doc:
            for block in page.get_text("dict")["blocks"]:
                if block.get("type") != 0:
                    continue
                for line in block["lines"]:
                    text = "".join(span["text"] for span in line["spans"]).strip()
                    if not text:
                        continue
                    x0 = min(span["bbox"][0] for span in line["spans"])
                    if re.match(r"^\d+\.\s", text) and x0 <= 58:
                        title = re.sub(r"^\d+\.\s+", "", text).strip()
                        decimal_titles.add(_norm_key(title))
                        decimal_titles.add(_norm_key(text))
                    elif x0 >= 62 and not re.match(r"^\d+\.\s", text):
                        bullet_keys.add(_norm_key(text))
    finally:
        doc.close()
    return decimal_titles, bullet_keys


def _is_decimal_heading(text: str, decimal_titles: set[str]) -> bool:
    text = text.strip()
    match = HEADING_RE.match(text)
    if match:
        return _norm_key(match.group(2)) in decimal_titles
    return _norm_key(text) in decimal_titles


def _matches_pdf_bullet(text: str, bullet_keys: set[str]) -> bool:
    key = _norm_key(text)
    if key in bullet_keys:
        return True
    prefix = key[:48]
    return any(b.startswith(prefix) or prefix.startswith(b[:48]) for b in bullet_keys if len(b) >= 20)


def strip_numbered_headings(doc: Document) -> None:
    """Keep numbered section titles intact (1. Problem Statement) — only strip SWOT duplicate."""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if text.startswith(SWOT_HEADING):
            continue
        # Do not strip main section numbers; PDF uses 1. 2. 3. for headings only.


def merge_wrapped_paragraphs(doc: Document) -> None:
    body = doc.element.body
    i = 0
    while i < len(body) - 1:
        a, b = body[i], body[i + 1]
        if a.tag != qn("w:p") or b.tag != qn("w:p"):
            i += 1
            continue

        ta = "".join(t.text or "" for t in a.iter(qn("w:t"))).strip()
        tb = "".join(t.text or "" for t in b.iter(qn("w:t"))).strip()
        if not ta or not tb:
            i += 1
            continue
        if HEADING_RE.match(ta) or HEADING_RE.match(tb):
            i += 1
            continue
        if ta.endswith((".", ":", "?", "!")):
            i += 1
            continue
        # Mid-sentence line wraps from pdf2docx (e.g. "...customer experience aur" / "control me rakho.")
        if (
            not HEADING_RE.match(tb)
            and (tb[0].islower() or tb.startswith(("control ", "relationships ", "me ", "ke ", "aur ")))
            and not ta.endswith((".", ":", "?", "!"))
        ):
            merged = f"{ta} {tb}".strip()
            paragraph = Paragraph(a, doc)
            paragraph.text = merged
            body.remove(b)
            continue
        if not CONTINUATION_START.match(tb) and not tb[0].islower():
            i += 1
            continue

        merged = f"{ta} {tb}".strip()
        paragraph = Paragraph(a, doc)
        paragraph.text = merged
        body.remove(b)
        continue


def parse_swot_block(text: str) -> dict[str, list[str]] | None:
    headers = ("Strengths", "Weaknesses", "Opportunities", "Threats")
    start = text.find("6. SWOT Analysis")
    if start < 0:
        return None
    end = text.find("Suggested response", start)
    if end < 0:
        end = text.find("7. Business Model", start)
    if end < 0:
        return None

    buckets: dict[str, list[str]] = {h: [] for h in headers}
    current: str | None = None
    for raw in text[start:end].splitlines():
        line = raw.strip()
        if not line or line == "6. SWOT Analysis":
            continue
        matched = next((h for h in headers if line == h), None)
        if matched:
            current = matched
            continue
        if current:
            buckets[current].append(line)
    return buckets if any(buckets.values()) else None


def parse_all_swot_blocks(pdf_path: str) -> list[dict[str, list[str]]]:
    doc = fitz.open(pdf_path)
    blocks: list[dict[str, list[str]]] = []
    for page in doc:
        parsed = parse_swot_block(page.get_text())
        if parsed:
            blocks.append(parsed)
    doc.close()
    return blocks


def normalize_swot_heading(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text or "Problem → Solution" in text:
            continue
        first_line = text.splitlines()[0].strip()
        match = HEADING_RE.match(first_line)
        if match:
            first_line = match.group(2).strip()
        if first_line == "SWOT Analysis" or first_line.startswith("SWOT Analysis "):
            paragraph.text = SWOT_HEADING


def _set_table_paragraph_style(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        p_style = OxmlElement("w:pStyle")
        p_pr.insert(0, p_style)
    p_style.set(qn("w:val"), "TableParagraph")


def _apply_cell_shading(cell, fill: str = SWOT_CELL_FILL) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)


def _clear_cell(cell) -> None:
    for paragraph in cell.paragraphs[1:]:
        element = paragraph._element
        element.getparent().remove(element)
    cell.paragraphs[0].clear()


def _fill_swot_cell(cell, header: str, items: list[str]) -> None:
    _clear_cell(cell)
    header_para = cell.paragraphs[0]
    header_run = header_para.add_run(header)
    header_run.bold = True
    _set_table_paragraph_style(header_para)

    for item in items:
        item_para = cell.add_paragraph()
        item_para.add_run(item)
        _set_table_paragraph_style(item_para)
        add_list_numbering(item_para, num_id=BULLET_NUM_ID)

    _apply_cell_shading(cell)


def _border_side(tag: str) -> OxmlElement:
    el = OxmlElement(f"w:{tag}")
    el.set(qn("w:val"), "single")
    el.set(qn("w:sz"), "6")
    el.set(qn("w:space"), "0")
    el.set(qn("w:color"), "CBD5E1")
    return el


def apply_swot_table_borders(table) -> None:
    """Visible box borders like Smallpdf (tblBorders, not style-only Table Grid)."""
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)

    for local in ("tblStyle", "tblBorders", "tblLayout", "jc", "tblInd"):
        existing = tbl_pr.find(qn(f"w:{local}"))
        if existing is not None and local != "tblInd":
            tbl_pr.remove(existing)

    borders = OxmlElement("w:tblBorders")
    for side in ("top", "left", "bottom", "right", "insideH", "insideV"):
        borders.append(_border_side(side))
    tbl_pr.append(borders)

    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "left")
    tbl_pr.append(jc)

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "42")
    tbl_ind.set(qn("w:type"), "dxa")


def fix_existing_swot_table_borders(doc: Document) -> None:
    for table in doc.tables:
        text = "".join(cell.text for row in table.rows for cell in row.cells)
        if (
            "Strengths" in text
            and "Weaknesses" in text
            and "Opportunities" in text
            and "Threats" in text
            and len(table.rows) == 2
            and len(table.columns) == 2
        ):
            apply_swot_table_borders(table)
            for row in table.rows:
                for cell in row.cells:
                    _apply_cell_shading(cell)


def ensure_swot_table(doc: Document, pdf_path: str) -> None:
    blocks = parse_all_swot_blocks(pdf_path)
    if not blocks:
        return

    body = doc.element.body
    block_idx = 0

    for child in list(body):
        if child.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
        if text not in (SWOT_HEADING, "SWOT Analysis"):
            continue
        if text == "SWOT Analysis":
            paragraph = Paragraph(child, doc)
            paragraph.text = SWOT_HEADING
        if block_idx >= len(blocks):
            break

        buckets = blocks[block_idx]
        block_idx += 1

        cursor = body.index(child) + 1
        while cursor < len(body):
            nxt = body[cursor]
            if nxt.tag == qn("w:p"):
                nt = "".join(t.text or "" for t in nxt.iter(qn("w:t"))).strip()
                if nt.startswith("Differentiate on trust") or nt.startswith("Business Model"):
                    break
                if nt.startswith("Suggested response"):
                    break
            if nxt.tag == qn("w:tbl"):
                tbl_text = "".join(t.text or "" for t in nxt.iter(qn("w:t")))
                if "Strengths" in tbl_text and "Weaknesses" in tbl_text:
                    break
            body.remove(nxt)

        tbl = doc.add_table(rows=2, cols=2)
        _fill_swot_cell(tbl.cell(0, 0), "Strengths", buckets["Strengths"])
        _fill_swot_cell(tbl.cell(0, 1), "Weaknesses", buckets["Weaknesses"])
        _fill_swot_cell(tbl.cell(1, 0), "Opportunities", buckets["Opportunities"])
        _fill_swot_cell(tbl.cell(1, 1), "Threats", buckets["Threats"])
        apply_swot_table_borders(tbl)

        tbl_el = tbl._tbl
        tbl_el.getparent().remove(tbl_el)
        body.insert(cursor, tbl_el)


def cleanup_swot_orphans(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
        if text != "Suggested response":
            continue

        cursor = body.index(child) + 1
        while cursor < len(body):
            nxt = body[cursor]
            if nxt.tag != qn("w:p"):
                break
            nt = "".join(t.text or "" for t in nxt.iter(qn("w:t"))).strip()
            if nt.startswith("Differentiate on trust") or nt.startswith("Business Model"):
                break
            body.remove(nxt)


def _paragraph_in_table(paragraph: Paragraph) -> bool:
    node = paragraph._element.getparent()
    while node is not None:
        if node.tag == qn("w:tbl"):
            return True
        node = node.getparent()
    return False


def _paragraph_in_ledger_table(paragraph: Paragraph) -> bool:
    node = paragraph._element.getparent()
    while node is not None:
        if node.tag == qn("w:tbl"):
            text = "".join(t.text or "" for t in node.iter(qn("w:t")))
            return any(m in text for m in LEDGER_MARKERS)
        node = node.getparent()
    return False


def strip_all_content_drawings(doc: Document) -> None:
    """Remove pdf2docx vector/shape artifacts; watermarks are injected later."""
    for paragraph in doc.element.body.iter(qn("w:p")):
        for run in list(paragraph.findall(qn("w:r"))):
            if run.find(qn("w:drawing")) is not None or run.find(qn("w:pict")) is not None:
                paragraph.remove(run)


def _clear_paragraph_indent(p_pr) -> None:
    ind = p_pr.find(qn("w:ind"))
    if ind is not None:
        p_pr.remove(ind)


def _set_hanging_indent(p_pr, left_twips: str, hanging_twips: str) -> None:
    _clear_paragraph_indent(p_pr)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), left_twips)
    ind.set(qn("w:hanging"), hanging_twips)
    p_pr.append(ind)


def _set_section_indent(p_pr) -> None:
    _clear_paragraph_indent(p_pr)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), SECTION_IND_LEFT)
    p_pr.append(ind)


def _set_run_font(run, half_points: str, bold: bool | None = None) -> None:
    r_pr = run._r.get_or_add_rPr()
    sz = r_pr.find(qn("w:sz"))
    if sz is None:
        sz = OxmlElement("w:sz")
        r_pr.append(sz)
    sz.set(qn("w:val"), half_points)
    sz_cs = r_pr.find(qn("w:szCs"))
    if sz_cs is None:
        sz_cs = OxmlElement("w:szCs")
        r_pr.append(sz_cs)
    sz_cs.set(qn("w:val"), half_points)
    if bold is not None:
        b = r_pr.find(qn("w:b"))
        if bold:
            if b is None:
                b = OxmlElement("w:b")
                r_pr.append(b)
        elif b is not None:
            r_pr.remove(b)


def _apply_section_heading_format(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    _set_section_indent(p_pr)
    for run in paragraph.runs:
        _set_run_font(run, HEADING_FONT_HALFPTS, bold=True)


def _apply_body_format(paragraph: Paragraph) -> None:
    for run in paragraph.runs:
        _set_run_font(run, BODY_FONT_HALFPTS, bold=False)


def add_list_numbering(paragraph: Paragraph, num_id: str = DECIMAL_NUM_ID) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:numPr")) is not None:
        return
    _clear_paragraph_indent(p_pr)
    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is not None:
        p_pr.remove(p_style)
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), num_id)
    num_pr.append(ilvl)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)
    if num_id == BULLET_NUM_ID:
        _set_hanging_indent(p_pr, BULLET_IND_LEFT, BULLET_IND_HANGING)
    _apply_body_format(paragraph)


def _clone_watermark_run() -> OxmlElement:
    import math

    from docx.oxml import parse_xml

    run_xml = WATERMARK_RUN_PATH.read_text(encoding="utf-8")
    x0, y0, x1, y1 = WATERMARK_BBOX
    width = round(WATERMARK_WIDTH_PT, 2)
    margin_left = round(WATERMARK_MARGIN_LEFT_PT, 2)
    top = round(WATERMARK_TOP_PT, 2)
    angle = math.degrees(math.atan2(y0 - y1, x1 - x0))
    rotation = int(round(360 + angle if angle < 0 else angle))
    run_xml = re.sub(
        r"margin-left:[^;]+;margin-top:[^;]+;width:[^;]+;height:[^;]+;",
        f"margin-left:{margin_left}pt;margin-top:{top}pt;width:{width}pt;height:48pt;",
        run_xml,
    )
    run_xml = re.sub(r"rotation:\d+", f"rotation:{rotation}", run_xml)
    wrapped = (
        '<w:wrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:v="urn:schemas-microsoft-com:vml" '
        'xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:w10="urn:schemas-microsoft-com:office:word">'
        f"{run_xml}</w:wrap>"
    )
    return deepcopy(parse_xml(wrapped)[0])


def _pt_to_emu(pt: float) -> int:
    return int(round(pt * 12700))


def _clone_watermark_drawing_run(doc_id: int) -> OxmlElement:
    from docx.oxml import parse_xml

    if not WATERMARK_DRAWING_PATH.is_file():
        return _clone_watermark_run()
    run_xml = WATERMARK_DRAWING_PATH.read_text(encoding="utf-8")
    pos_h = _pt_to_emu(WATERMARK_MARGIN_LEFT_PT)
    pos_v = _pt_to_emu(WATERMARK_TOP_PT)
    width_emu = _pt_to_emu(WATERMARK_WIDTH_PT)
    run_xml = re.sub(
        r"(<wp:positionH[^>]*>\s*<wp:posOffset>)\d+(</wp:posOffset>)",
        rf"\g<1>{pos_h}\2",
        run_xml,
        count=1,
    )
    run_xml = re.sub(
        r"(<wp:positionV[^>]*>\s*<wp:posOffset>)\d+(</wp:posOffset>)",
        rf"\g<1>{pos_v}\2",
        run_xml,
        count=1,
    )
    run_xml = re.sub(r'cx="\d+"', f'cx="{width_emu}"', run_xml)
    run_xml = re.sub(r'wp:docPr id="\d+"', f'wp:docPr id="{doc_id}"', run_xml)
    run_xml = re.sub(r'wps:cNvPr id="\d+"', f'wps:cNvPr id="{doc_id}"', run_xml)
    wrapped = (
        '<w:wrap xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" '
        'xmlns:wp="http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing" '
        'xmlns:a="http://schemas.openxmlformats.org/drawingml/2006/main" '
        'xmlns:wps="http://schemas.microsoft.com/office/word/2010/wordprocessingShape">'
        f"{run_xml}</w:wrap>"
    )
    return deepcopy(parse_xml(wrapped)[0])


def extract_pdf_watermark_markers(pdf_path: str, max_pages: int = MAX_WATERMARKS) -> list[str]:
    """First heading line per PDF page that contains CONFIDENTIAL watermark."""
    markers: list[str] = []
    doc = fitz.open(pdf_path)
    try:
        for page in doc[:max_pages]:
            if "CONFIDENTIAL" not in page.get_text():
                continue
            tops: list[tuple[float, float, str]] = []
            for block in page.get_text("dict")["blocks"]:
                if block.get("type") != 0:
                    continue
                for line in block["lines"]:
                    text = "".join(span["text"] for span in line["spans"]).strip()
                    if not text or "CONFIDENTIAL" in text:
                        continue
                    y0 = line["bbox"][1]
                    x0 = min(span["bbox"][0] for span in line["spans"])
                    tops.append((y0, x0, text))
            tops.sort()
            marker: str | None = None
            if not tops:
                continue
            # Page 1: anchor at document title so watermark sits over full page
            if len(markers) == 0:
                marker = tops[0][2]
            else:
                for _y, x0, text in tops:
                    if re.match(r"^\d+\.\s", text) and x0 <= 55:
                        marker = text
                        break
                if marker is None:
                    marker = tops[0][2]
            if marker:
                markers.append(marker)
    finally:
        doc.close()
    return markers


def _paragraph_matches_marker(text: str, marker: str) -> bool:
    text_key = _norm_key(text)
    marker_key = _norm_key(marker)
    if text_key == marker_key:
        return True
    if text_key.startswith(marker_key[:32]):
        return True
    return marker_key[:32] in text_key[: max(len(text_key), 32)]


def _insert_watermark_before(target, run) -> None:
    new_p = OxmlElement("w:p")
    p_pr = OxmlElement("w:pPr")
    spacing = OxmlElement("w:spacing")
    spacing.set(qn("w:before"), "0")
    spacing.set(qn("w:after"), "0")
    spacing.set(qn("w:line"), "0")
    spacing.set(qn("w:lineRule"), "exact")
    p_pr.append(spacing)
    new_p.append(p_pr)
    new_p.append(deepcopy(run))
    target.addprevious(new_p)


def _make_pg_mar() -> OxmlElement:
    pg_mar = OxmlElement("w:pgMar")
    for key, val in (
        ("top", "578"),
        ("right", "1316"),
        ("bottom", "1440"),
        ("left", "1020"),
        ("header", "720"),
        ("footer", "720"),
        ("gutter", "0"),
    ):
        pg_mar.set(qn(f"w:{key}"), val)
    return pg_mar


def _make_sect_pr(landscape: bool = False) -> OxmlElement:
    sect = OxmlElement("w:sectPr")
    pg_sz = OxmlElement("w:pgSz")
    if landscape:
        pg_sz.set(qn("w:w"), "16838")
        pg_sz.set(qn("w:h"), "11906")
        pg_sz.set(qn("w:orient"), "landscape")
    else:
        pg_sz.set(qn("w:w"), "11906")
        pg_sz.set(qn("w:h"), "16838")
    sect.append(pg_sz)
    sect.append(_make_pg_mar())
    return sect


def _set_paragraph_sect_pr(paragraph: Paragraph, sect_pr: OxmlElement) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    existing = p_pr.find(qn("w:sectPr"))
    if existing is not None:
        p_pr.remove(existing)
    p_pr.append(deepcopy(sect_pr))


def _is_ledger_table(table) -> bool:
    text = "".join(cell.text for row in table.rows for cell in row.cells)
    return any(marker in text for marker in LEDGER_MARKERS) or len(table.columns) >= 6


def _scale_table_to_width(table, max_twips: int) -> None:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.insert(0, tbl_w)
    tbl_w.set(qn("w:type"), "pct")
    tbl_w.set(qn("w:w"), "5000")

    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")

    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        return
    cols = grid.findall(qn("w:gridCol"))
    total = sum(int(col.get(qn("w:w"), "0")) for col in cols)
    if total <= 0:
        return
    if total > max_twips:
        scale = max_twips / total
        for col in cols:
            width = int(col.get(qn("w:w"), "0"))
            col.set(qn("w:w"), str(max(1, int(width * scale))))


def fit_all_tables_to_page(doc: Document) -> None:
    """Scale tables to printable width; wide ledger tables use landscape width."""
    for table in doc.tables:
        max_w = LANDSCAPE_CONTENT_TWIPS if _is_ledger_table(table) else PORTRAIT_CONTENT_TWIPS
        _scale_table_to_width(table, max_w)


def apply_landscape_ledger_sections(doc: Document) -> None:
    """PDF pages 9–19 are landscape ledger tables — match with a landscape section."""
    start_idx: int | None = None
    for i, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if text.startswith(LEDGER_START_PREFIX) and "Part 2" not in text:
            start_idx = i
            break
    if start_idx is None or start_idx == 0:
        return

    end_idx: int | None = None
    nav_count = 0
    for i, paragraph in enumerate(doc.paragraphs):
        if i <= start_idx:
            continue
        text = paragraph.text.strip()
        if "Problem → Solution" in text and "Co-founder KRA/KPI" in text:
            nav_count += 1
            if nav_count >= 2:
                end_idx = i - 1
                break
        if text == "Co-founder KRA/KPI" and i > start_idx + 3:
            end_idx = i - 1
            break

    if end_idx is None or end_idx < start_idx:
        end_idx = min(len(doc.paragraphs) - 1, start_idx + 40)

    _set_paragraph_sect_pr(doc.paragraphs[start_idx - 1], _make_sect_pr(landscape=False))
    _set_paragraph_sect_pr(doc.paragraphs[end_idx], _make_sect_pr(landscape=True))

    body_paras = [c for c in doc.element.body if c.tag == qn("w:p")]
    if body_paras:
        _set_paragraph_sect_pr(Paragraph(body_paras[-1], doc), _make_sect_pr(landscape=False))


def clear_table_cell_fills(doc: Document) -> None:
    """Remove opaque table cell fills so header watermark shows through (like PDF)."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                tc_pr = cell._tc.get_or_add_tcPr()
                shd = tc_pr.find(qn("w:shd"))
                if shd is not None:
                    tc_pr.remove(shd)
                clear = OxmlElement("w:shd")
                clear.set(qn("w:val"), "clear")
                clear.set(qn("w:color"), "auto")
                clear.set(qn("w:fill"), "auto")
                tc_pr.append(clear)


def remove_body_watermarks(doc: Document) -> None:
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
        if text:
            for run in list(child.findall(qn("w:r"))):
                if _is_confidential_watermark_run(run):
                    child.remove(run)
            continue
        if any(_is_confidential_watermark_run(run) for run in child.findall(qn("w:r"))):
            body.remove(child)


def inject_watermark_header(doc: Document) -> None:
    """Single header watermark on the first section only (PDF pages 1–8)."""
    if not WATERMARK_RUN_PATH.is_file() or not doc.sections:
        return
    for si, section in enumerate(doc.sections):
        section.header.is_linked_to_previous = False
        if si == 0:
            header = section.header
            if header.paragraphs:
                hp = header.paragraphs[0]
                hp.clear()
            else:
                hp = header.add_paragraph()
            p_pr = hp._p.get_or_add_pPr()
            p_style = p_pr.find(qn("w:pStyle"))
            if p_style is not None:
                p_pr.remove(p_style)
            spacing = p_pr.find(qn("w:spacing"))
            if spacing is None:
                spacing = OxmlElement("w:spacing")
                p_pr.append(spacing)
            spacing.set(qn("w:before"), "0")
            spacing.set(qn("w:after"), "0")
            spacing.set(qn("w:line"), "0")
            spacing.set(qn("w:lineRule"), "exact")
            hp._p.append(_clone_watermark_run())
        elif section.header.paragraphs:
            section.header.paragraphs[0].clear()


def clear_header_watermarks(doc: Document) -> None:
    for section in doc.sections:
        section.header.is_linked_to_previous = False
        for paragraph in section.header.paragraphs:
            paragraph.clear()


def inject_body_page_watermarks(doc: Document, pdf_path: str) -> None:
    """Page-anchored body watermarks (Smallpdf-style) — continuous through tables."""
    markers = extract_pdf_watermark_markers(pdf_path)
    if len(markers) < MAX_WATERMARKS:
        markers = list(WATERMARK_ANCHORS[:MAX_WATERMARKS])
    used_indices: set[int] = set()
    doc_id = 100
    placed = 0
    for marker in markers:
        target_p = None
        for i, paragraph in enumerate(doc.paragraphs):
            if i in used_indices:
                continue
            if _paragraph_matches_marker(paragraph.text, marker):
                target_p = paragraph._p
                used_indices.add(i)
                break
        if target_p is None:
            continue
        doc_id += 1
        _insert_watermark_before(target_p, _clone_watermark_drawing_run(doc_id))
        placed += 1
    if placed == 0:
        inject_watermark_header(doc)


def inject_confidential_watermarks(doc: Document, pdf_path: str) -> None:
    remove_body_watermarks(doc)
    clear_header_watermarks(doc)
    inject_body_page_watermarks(doc, pdf_path)


def _is_confidential_watermark_run(run) -> bool:
    raw = run.xml
    if "CONFIDENTIAL" not in raw:
        return False
    return "v:shape" in raw or "w:pict" in raw or "w:drawing" in raw


def cleanup_empty_list_paragraphs(doc: Document) -> None:
    """Remove empty paragraphs that only show a stray bullet dot."""
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:p"):
            continue
        text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
        if text:
            continue
        p_pr = child.find(qn("w:pPr"))
        has_num = p_pr is not None and p_pr.find(qn("w:numPr")) is not None
        has_wm = any(_is_confidential_watermark_run(run) for run in child.findall(qn("w:r")))
        if has_num and not has_wm:
            body.remove(child)
            continue
        if p_pr is not None and has_num:
            num_pr = p_pr.find(qn("w:numPr"))
            if num_pr is not None:
                p_pr.remove(num_pr)


def _paragraph_has_confidential_watermark(paragraph: Paragraph) -> bool:
    return any(_is_confidential_watermark_run(run) for run in paragraph._p.findall(qn("w:r")))


def strip_drawings_from_numbered_paragraphs(doc: Document) -> None:
    """Remove pdf2docx bullet drawings where Word list numbering is used (fixes ::)."""
    for paragraph in doc.element.body.iter(qn("w:p")):
        p_pr = paragraph.find(qn("w:pPr"))
        if p_pr is None or p_pr.find(qn("w:numPr")) is None:
            continue
        for run in list(paragraph.findall(qn("w:r"))):
            if run.find(qn("w:drawing")) is not None and not _is_confidential_watermark_run(run):
                paragraph.remove(run)


def _element_in_swot_table(element) -> bool:
    node = element
    while node is not None:
        if node.tag == qn("w:tbl"):
            text = "".join(t.text or "" for t in node.iter(qn("w:t")))
            return (
                "Strengths" in text
                and "Weaknesses" in text
                and "Opportunities" in text
                and "Threats" in text
            )
        node = node.getparent()
    return False


def strip_orphan_bullet_drawings(doc: Document) -> None:
    """Remove pdf2docx-only bullet shapes that show as extra black dots in Word."""
    body = doc.element.body
    for child in list(body):
        if child.tag != qn("w:p"):
            continue
        if _element_in_swot_table(child):
            continue

        text = "".join(t.text or "" for t in child.iter(qn("w:t"))).strip()
        runs_with_drawings = [
            run
            for run in child.findall(qn("w:r"))
            if run.find(qn("w:drawing")) is not None
        ]
        if not runs_with_drawings:
            continue

        if any(_is_confidential_watermark_run(run) for run in runs_with_drawings):
            continue

        p_pr = child.find(qn("w:pPr"))
        has_num = p_pr is not None and p_pr.find(qn("w:numPr")) is not None

        if not text:
            body.remove(child)
            continue

        if has_num or len(text) < 48:
            for run in runs_with_drawings:
                child.remove(run)


NO_BULLET_EXACT = {
    "Formats",
    "Suggested response",
    "Core marketing message",
    "Recommended product architecture",
    "Recommended pricing",
    "Sales funnel",
    "Financial structure",
    "Base-case targets",
    "Revenue streams",
    "Gap ka simple version:",
    "Existing market bands:",
    "Current shelf reality:",
    "Recommended price gap for The786: accessible premium band.",
    "90-day GTM priorities",
    "Suggested KPI scorecard weight",
    "Strategy Role Priority",
    "Segment Product Role Launch Phase",
    "SKU MRP (incl. GST) Role",
    "KRA What they own Core KPIs",
    "Strengths",
    "Weaknesses",
    "Opportunities",
    "Threats",
    SWOT_HEADING,
    "SWOT Analysis",
}

NO_BULLET_PREFIXES = (
    "Simple answer:",
    "Problem →",
    "The786 —",
    "Brand Position",
    "Primary Model",
    "Primary Engine",
    "Founder note:",
    "Old working Sheet",
    "Slide ",
    "Page ",
)

NO_BULLET_RE = (
    re.compile(r"^Slide \d+"),
    re.compile(r"^Page \d+"),
)


def should_have_bullet(text: str) -> bool:
    text = text.strip()
    if not text or text in NO_BULLET_EXACT:
        return False
    if text.startswith(NO_BULLET_PREFIXES):
        return False
    if any(pattern.match(text) for pattern in NO_BULLET_RE):
        return False
    if text.endswith(":") and text in {
        "Gap ka simple version:",
        "Existing market bands:",
        "Current shelf reality:",
    }:
        return False
    return True


def remove_list_numbering(paragraph: Paragraph) -> None:
    p_pr = paragraph._p.find(qn("w:pPr"))
    if p_pr is None:
        return
    num_pr = p_pr.find(qn("w:numPr"))
    if num_pr is not None:
        p_pr.remove(num_pr)


SKIP_LIST = NO_BULLET_EXACT


def normalize_plain_paragraphs(doc: Document) -> None:
    """Clear pdf2docx firstLine=0 on plain paragraphs (breaks visual alignment)."""
    for paragraph in doc.paragraphs:
        if not paragraph.text.strip():
            continue
        p_pr = paragraph._p.find(qn("w:pPr"))
        if p_pr is None or p_pr.find(qn("w:numPr")) is not None:
            continue
        ind = p_pr.find(qn("w:ind"))
        if ind is not None and ind.get(qn("w:hanging")) is None:
            if ind.get(qn("w:firstLine")) == "0" or ind.get(qn("w:left")) in {"60", "0"}:
                p_pr.remove(ind)


def apply_native_lists(doc: Document, pdf_path: str) -> None:
    decimal_titles, bullet_keys = build_pdf_list_map(pdf_path)
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if _paragraph_in_table(paragraph) or _paragraph_in_ledger_table(paragraph):
            continue
        if _is_decimal_heading(text, decimal_titles):
            remove_list_numbering(paragraph)
            _apply_section_heading_format(paragraph)
        elif should_have_bullet(text) or _matches_pdf_bullet(text, bullet_keys):
            add_list_numbering(paragraph, num_id=BULLET_NUM_ID)
        else:
            remove_list_numbering(paragraph)

    for paragraph in doc.paragraphs:
        p_pr = paragraph._p.find(qn("w:pPr"))
        if p_pr is None:
            continue
        num_pr = p_pr.find(qn("w:numPr"))
        if num_pr is None:
            continue
        num_id_el = num_pr.find(qn("w:numId"))
        if num_id_el is not None and num_id_el.get(qn("w:val")) == BULLET_NUM_ID:
            _set_hanging_indent(p_pr, BULLET_IND_LEFT, BULLET_IND_HANGING)
            p_style = p_pr.find(qn("w:pStyle"))
            if p_style is not None:
                p_pr.remove(p_style)
            _apply_body_format(paragraph)


def normalize_plain_paragraphs(doc: Document) -> None:
    """Clear pdf2docx firstLine=0 indents on non-list body paragraphs."""
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        p_pr = paragraph._p.find(qn("w:pPr"))
        if p_pr is None:
            continue
        if p_pr.find(qn("w:numPr")) is not None:
            continue
        if _is_decimal_heading(text, set()) and HEADING_RE.match(text):
            continue
        ind = p_pr.find(qn("w:ind"))
        if ind is None:
            continue
        if ind.get(qn("w:hanging")) is None and ind.get(qn("w:firstLine")) == "0":
            p_pr.remove(ind)


def inject_template_styles(docx_path: str) -> None:
    if not TEMPLATE_DIR.exists():
        return
    import io

    with zipfile.ZipFile(docx_path, "r") as zin:
        entries = {name: zin.read(name) for name in zin.namelist() if not name.endswith("/")}

    for fname in ("numbering.xml", "styles.xml", "fontTable.xml"):
        src = TEMPLATE_DIR / fname
        if src.is_file():
            entries[f"word/{fname}"] = src.read_bytes()

    numbering_rels = TEMPLATE_DIR / "_rels" / "numbering.xml.rels"
    if numbering_rels.is_file():
        entries["word/_rels/numbering.xml.rels"] = numbering_rels.read_bytes()

    bullet_img = TEMPLATE_DIR / "media" / "image1.png"
    if bullet_img.is_file():
        entries["word/media/image1.png"] = bullet_img.read_bytes()

    ct = entries["[Content_Types].xml"].decode("utf-8", errors="ignore")
    if "numbering.xml" not in ct:
        ct = ct.replace(
            "</Types>",
            '  <Override PartName="/word/numbering.xml" '
            'ContentType="application/vnd.openxmlformats-officedocument'
            '.wordprocessingml.numbering+xml"/>\n</Types>',
        )
    if bullet_img.is_file() and "/word/media/image1.png" not in ct:
        ct = ct.replace(
            "</Types>",
            '  <Override PartName="/word/media/image1.png" '
            'ContentType="image/png"/>\n</Types>',
        )
    entries["[Content_Types].xml"] = ct.encode("utf-8")

    rels_key = "word/_rels/document.xml.rels"
    rels = entries[rels_key].decode("utf-8", errors="ignore")
    if "numbering.xml" not in rels:
        nums = [int(n) for n in re.findall(r'Id="rId(\d+)"', rels)]
        next_id = max(nums, default=0) + 1
        rels = rels.replace(
            "</Relationships>",
            f'  <Relationship Id="rId{next_id}" '
            f'Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/numbering" '
            f'Target="numbering.xml"/>\n</Relationships>',
        )
        entries[rels_key] = rels.encode("utf-8")

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in entries.items():
            zout.writestr(name, data)
    Path(docx_path).write_bytes(buffer.getvalue())


def _referenced_media(entries: dict[str, bytes]) -> set[str]:
    refs: set[str] = set()
    for name, data in entries.items():
        if not (name.endswith(".xml") or name.endswith(".rels")):
            continue
        text = data.decode("utf-8", errors="ignore")
        for match in re.findall(r"media/([^\"'>\s]+)", text):
            refs.add(match)
    return refs


def strip_embedded_images(docx_path: str) -> None:
    """Remove inline raster images from body; keep numbering bullet asset only."""
    import io

    with zipfile.ZipFile(docx_path, "r") as zin:
        entries = {name: zin.read(name) for name in zin.namelist() if not name.endswith("/")}

    doc_xml = entries["word/document.xml"].decode("utf-8", errors="ignore")

    def _keep_watermark_drawing(match: re.Match[str]) -> str:
        block = match.group(0)
        if "CONFIDENTIAL" in block and "wp:anchor" in block:
            return block
        return ""

    doc_xml = re.sub(r"<w:drawing>.*?</w:drawing>", _keep_watermark_drawing, doc_xml, flags=re.DOTALL)
    doc_xml = re.sub(
        r"<w:pict>(?!.*CONFIDENTIAL).*?</w:pict>",
        "",
        doc_xml,
        flags=re.DOTALL,
    )
    entries["word/document.xml"] = doc_xml.encode("utf-8")

    rels_key = "word/_rels/document.xml.rels"
    if rels_key in entries:
        rels = entries[rels_key].decode("utf-8", errors="ignore")
        rels = re.sub(
            r'<Relationship[^>]+Type="[^"]*/image"[^>]*/>\s*',
            "",
            rels,
            flags=re.IGNORECASE,
        )
        entries[rels_key] = rels.encode("utf-8")

    keep_media = {"image1.png"}
    for name in list(entries):
        if name.startswith("word/media/") and name.split("/")[-1] not in keep_media:
            del entries[name]

    ct = entries["[Content_Types].xml"].decode("utf-8", errors="ignore")
    ct = re.sub(
        r'\s*<Override PartName="/word/media/image(?!1\.png)[^"]+"[^>]*/>',
        "",
        ct,
        flags=re.IGNORECASE,
    )
    entries["[Content_Types].xml"] = ct.encode("utf-8")

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in entries.items():
            zout.writestr(name, data)
    Path(docx_path).write_bytes(buffer.getvalue())


def compress_docx_package(docx_path: str) -> None:
    """Drop unreferenced media and recompress images to reduce file size."""
    import hashlib
    import io

    with zipfile.ZipFile(docx_path, "r") as zin:
        entries = {name: zin.read(name) for name in zin.namelist() if not name.endswith("/")}

    refs = _referenced_media(entries)
    for name in list(entries):
        if name.startswith("word/media/") and name.split("/")[-1] not in refs:
            del entries[name]

    try:
        from PIL import Image
    except ImportError:
        Image = None

    if Image is not None:
        for name in list(entries):
            if not name.startswith("word/media/"):
                continue
            lower = name.lower()
            if not (lower.endswith(".png") or lower.endswith(".jpg") or lower.endswith(".jpeg")):
                continue
            try:
                img = Image.open(io.BytesIO(entries[name]))
                buf = io.BytesIO()
                if lower.endswith(".png"):
                    img.save(buf, format="PNG", optimize=True)
                else:
                    rgb = img.convert("RGB")
                    quality = 50 if len(entries[name]) > 50_000 else 72
                    rgb.save(buf, format="JPEG", quality=quality, optimize=True)
                entries[name] = buf.getvalue()
            except OSError:
                continue

    # Deduplicate identical media blobs
    hash_to_name: dict[str, str] = {}
    rename: dict[str, str] = {}
    for name in sorted(entries):
        if not name.startswith("word/media/"):
            continue
        digest = hashlib.sha256(entries[name]).hexdigest()
        if digest in hash_to_name:
            rename[name.split("/")[-1]] = hash_to_name[digest].split("/")[-1]
        else:
            hash_to_name[digest] = name

    if rename:
        for part in list(entries):
            if not (part.endswith(".xml") or part.endswith(".rels")):
                continue
            text = entries[part].decode("utf-8", errors="ignore")
            for old, new in rename.items():
                text = text.replace(f"media/{old}", f"media/{new}")
            entries[part] = text.encode("utf-8")
        for name in list(entries):
            if name.startswith("word/media/") and name.split("/")[-1] in rename:
                del entries[name]

    buffer = io.BytesIO()
    with zipfile.ZipFile(buffer, "w", compression=zipfile.ZIP_DEFLATED) as zout:
        for name, data in entries.items():
            zout.writestr(name, data)
    Path(docx_path).write_bytes(buffer.getvalue())


def merge_nav_line(doc: Document) -> None:
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip() != "Co-founder KRA/KPI":
            continue
        if i == 0:
            continue
        prev = doc.paragraphs[i - 1]
        if "Problem → Solution" in prev.text and "Co-founder" not in prev.text:
            prev.text = prev.text.rstrip() + " Co-founder KRA/KPI"
            p = paragraph._element
            parent = p.getparent()
            if parent is not None:
                parent.remove(p)
        break


def transform_to_smallpdf(docx_path: str, pdf_path: str) -> None:
    doc = Document(docx_path)
    strip_all_content_drawings(doc)
    merge_wrapped_paragraphs(doc)
    merge_nav_line(doc)
    strip_numbered_headings(doc)
    normalize_swot_heading(doc)
    ensure_swot_table(doc, pdf_path)
    cleanup_swot_orphans(doc)
    fix_existing_swot_table_borders(doc)
    doc.save(docx_path)
    inject_template_styles(docx_path)

    doc = Document(docx_path)
    normalize_plain_paragraphs(doc)
    apply_native_lists(doc, pdf_path)
    strip_drawings_from_numbered_paragraphs(doc)
    strip_orphan_bullet_drawings(doc)
    cleanup_empty_list_paragraphs(doc)
    apply_landscape_ledger_sections(doc)
    fit_all_tables_to_page(doc)
    clear_table_cell_fills(doc)
    inject_confidential_watermarks(doc, pdf_path)
    doc.save(docx_path)
    strip_embedded_images(docx_path)
    compress_docx_package(docx_path)
