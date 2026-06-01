"""Rebuild pdf2docx output to match Smallpdf document structure."""
from __future__ import annotations

import re
from typing import Callable

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.table import Table
from docx.text.paragraph import Paragraph

HEADING_RE = re.compile(r"^\d+\.\s+\S")

PARAGRAPH_SPLIT_MARKERS = (
    "Consumer pain points:",
    "Body-related risks:",
    "Founder-side constraints:",
    "Brand challenge:",
    "Mass:",
    "Mid-range:",
    "Premium & luxury:",
    "Premium &amp; luxury:",
    "Current shelf reality:",
    "Recommended price gap for The786:",
    "2ml sample:",
    "10ml travel:",
    "50ml hero bottle:",
    "100ml hero bottle:",
    "Simple answer:",
    "Secondary sales channel:",
    "Later channels:",
    "Phase 2 extension:",
)

LEDGER_MARKERS = ("Sl. No.", "Debit(-)", "Credit(+)")


def _pdf_full_text(pdf_path: str) -> str:
    import fitz

    doc = fitz.open(pdf_path)
    text = "".join(page.get_text() for page in doc)
    doc.close()
    return text


def _pdf_block(full: str, start: str, end: str) -> list[str]:
    s = full.find(start)
    if s < 0:
        return []
    e = full.find(end, s + len(start))
    if e < 0:
        e = len(full)
    lines: list[str] = []
    for raw in full[s:e].splitlines():
        line = raw.strip()
        if not line or line == "CONFIDENTIAL":
            continue
        lines.append(line)
    return lines


def _para_text(element) -> str:
    return "".join(node.text or "" for node in element.iter(qn("w:t"))).strip()


def _replace_range_with_lines(
    doc: Document, start_marker: str, end_marker: str, lines: list[str]
) -> int:
    if not lines:
        return 0

    body = doc.element.body
    changed = 0
    search_from = 0

    while search_from < len(body):
        start_idx = None
        end_idx = None

        for i in range(search_from, len(body)):
            child = body[i]
            if child.tag != qn("w:p"):
                continue
            text = _para_text(child)
            if start_marker in text:
                start_idx = i
                break

        if start_idx is None:
            break

        for i in range(start_idx + 1, len(body)):
            child = body[i]
            if child.tag == qn("w:p"):
                text = _para_text(child)
                if end_marker in text:
                    end_idx = i
                    break
            elif child.tag == qn("w:tbl") and end_marker.startswith("Suggested"):
                end_idx = i
                break

        if end_idx is None:
            break

        for i in range(end_idx - 1, start_idx, -1):
            body.remove(body[i])

        insert_at = start_idx + 1
        for line in lines[1:]:
            p = doc.add_paragraph(line)
            body.insert(insert_at, p._element)
            insert_at += 1

        changed += 1
        search_from = insert_at + 1

    return changed


def _transform_sales(lines: list[str]) -> list[str]:
    out: list[str] = []
    i = 0
    while i < len(lines):
        if (
            i + 1 < len(lines)
            and lines[i] == "Discovery Set or 10ml trial"
            and lines[i + 1] == "50ml hero conversion"
        ):
            out.append("Discovery Set or 10ml trial 50ml hero conversion")
            i += 2
            continue
        out.append(lines[i])
        i += 1
    return out


def _transform_price(lines: list[str]) -> list[str]:
    out: list[str] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line.startswith("Simple answer:") and i + 1 < len(lines):
            out.append(f"{line} {lines[i + 1]}".strip())
            i += 2
            continue
        if line.endswith("trust-led") and i + 1 < len(lines):
            out.append(f"{line} {lines[i + 1]}".strip())
            i += 2
            continue
        out.append(line)
        i += 1
    return out


def _transform_financial(lines: list[str]) -> list[str]:
    if not lines:
        return lines

    out: list[str] = [lines[0]]
    i = 1
    while i < len(lines):
        line = lines[i]

        if line == "Recommended pricing":
            out.append(line)
            i += 1
            if i < len(lines) and lines[i] == "SKU":
                out.append("SKU MRP (incl. GST) Role")
                i += 3
            continue

        if (
            i + 2 < len(lines)
            and not lines[i + 1].startswith("₹")
            and lines[i + 1] in {"₹299", "₹799", "₹2,490", "₹1,299"}
        ):
            pass

        if i + 2 < len(lines) and lines[i + 1].startswith("₹"):
            out.append(f"{lines[i]} {lines[i+1]} {lines[i+2]}")
            i += 3
            continue

        if line in {"SKU", "MRP (incl. GST)", "Role"}:
            i += 1
            continue

        out.append(line)
        i += 1

    return out


def _transform_products(lines: list[str]) -> list[str]:
    if not lines:
        return lines

    out: list[str] = [lines[0]]
    i = 1
    while i < len(lines):
        line = lines[i]

        if line == "Recommended product architecture":
            out.append(line)
            i += 1
            if i < len(lines) and lines[i] == "Segment":
                out.append("Segment Product Role Launch Phase")
                i += 4
            continue

        if (
            i + 3 < len(lines)
            and lines[i + 3].startswith("Phase ")
            and lines[i + 1][0].isupper()
        ):
            out.append(f"{lines[i]} {lines[i+1]} {lines[i+2]} {lines[i+3]}")
            i += 4
            continue

        if line == "Formats":
            out.append(line)
            i += 1
            chunks: list[str] = []
            while i < len(lines) and not lines[i].startswith("Phase 2 extension"):
                chunks.append(lines[i])
                i += 1
            if chunks:
                out.append(" ".join(chunks))
            continue

        if line.startswith("Phase 2 extension"):
            out.append(line)
            i += 1
            continue

        out.append(line)
        i += 1

    return out


def rebuild_sections_from_pdf(doc: Document, pdf_path: str) -> None:
    full = _pdf_full_text(pdf_path)
    sections: list[tuple[str, str, Callable[[list[str]], list[str]]]] = [
        ("9. Sales Model", "10. Financial Modeling", _transform_sales),
        ("10. Financial Modeling", "11. GTM", _transform_financial),
        ("3. Product", "4. Market Gap", _transform_products),
        ("5. Price Market Gap", "6. SWOT Analysis", _transform_price),
    ]

    for start, end, transform in sections:
        block = transform(_pdf_block(full, start, end))
        _replace_range_with_lines(doc, start, end, block)


def _cell_text(cell) -> str:
    return " ".join(p.text.strip() for p in cell.paragraphs if p.text.strip())


def _row_texts(row) -> list[str]:
    seen: set[int] = set()
    texts: list[str] = []
    for cell in row.cells:
        cell_id = id(cell._tc)
        if cell_id in seen:
            continue
        seen.add(cell_id)
        text = _cell_text(cell)
        if text:
            texts.append(text)
    return texts


def _header_key(table: Table) -> str:
    if not table.rows:
        return ""
    return " | ".join(_row_texts(table.rows[0])).lower()


def _split_cell_two_rows(text: str) -> tuple[str, str]:
    if len(text) <= 42:
        return text, ""
    words = text.split()
    if len(words) <= 4:
        return text, ""
    mid = max(1, len(words) // 2)
    return " ".join(words[:mid]), " ".join(words[mid:])


def split_kra_tables(doc: Document, pdf_path: str) -> None:
    del pdf_path  # table row text already extracted by pdf2docx

    for table in list(doc.tables):
        key = _header_key(table)
        if "kra" not in key or "what they own" not in key:
            continue
        if len(table.rows) < 2:
            continue

        parent = table._tbl.getparent()
        if parent is None:
            continue
        idx = parent.index(table._tbl)
        parent.remove(table._tbl)

        for row in table.rows[1:]:
            values = (_row_texts(row) + ["", "", ""])[:3]
            if not any(values):
                continue

            mini = doc.add_table(rows=2, cols=3)
            mini.style = "Table Grid"
            for col, value in enumerate(values):
                top, bottom = _split_cell_two_rows(value)
                mini.cell(0, col).text = top
                mini.cell(1, col).text = bottom

            parent.insert(idx, mini._tbl)
            idx += 1


def _split_price_lines(parts: list[str]) -> list[str]:
    out: list[str] = []
    for part in parts:
        if part.startswith("Simple answer:"):
            out.append(part)
            continue
        chunks = re.split(r"(?<=[a-z])(?=₹5,600\+)", part)
        for chunk in chunks:
            out.extend(
                re.split(r"(?<=[a-z])(?=Recommended price gap for The786:)", chunk)
            )
    return [p.strip() for p in out if p.strip()]


def split_merged_paragraphs(doc: Document) -> None:
    for paragraph in list(doc.paragraphs):
        if not paragraph.text.strip():
            continue
        if HEADING_RE.match(paragraph.text.strip()):
            continue

        parts = [paragraph.text.strip()]
        for marker in PARAGRAPH_SPLIT_MARKERS:
            next_parts: list[str] = []
            for part in parts:
                if marker not in part:
                    next_parts.append(part)
                    continue
                idx = part.find(marker)
                if idx > 0:
                    next_parts.append(part[:idx].strip())
                next_parts.append(part[idx:].strip())
            parts = [p for p in next_parts if p.strip()]

        parts = _split_price_lines(parts)

        if len(parts) <= 1:
            continue

        parent = paragraph._element.getparent()
        if parent is None:
            continue
        idx = parent.index(paragraph._element)
        parent.remove(paragraph._element)
        for offset, part in enumerate(parts):
            p = doc.add_paragraph(part)
            parent.insert(idx + offset, p._element)


def flatten_table_to_lines(table: Table, *, skip_header: bool = False) -> list[str]:
    rows = [_row_texts(row) for row in table.rows]
    if not rows:
        return []

    lines: list[str] = []
    start = 1 if skip_header else 0
    if not skip_header and rows:
        lines.append(" ".join(rows[0]))

    for row in rows[start:]:
        line = " ".join(row).strip()
        if line:
            lines.append(line)
    return lines


def _replace_element_with_paragraphs(doc: Document, element, lines: list[str]) -> None:
    parent = element.getparent()
    if parent is None:
        return
    idx = parent.index(element)
    parent.remove(element)
    for offset, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph(line)
        parent.insert(idx + offset, p._element)


def flatten_text_layout_tables(doc: Document) -> None:
    for table in list(doc.tables):
        key = _header_key(table)
        if not key:
            continue

        flatten: list[str] | None = None

        if "strategy" in key and "role" in key and "priority" in key:
            flatten = flatten_table_to_lines(table, skip_header=False)
        elif key.replace(" ", "") in {"strategyrolepriority"}:
            flatten = flatten_table_to_lines(table, skip_header=False)
        elif "segment" in key and "product" in key and "launch phase" in key:
            flatten = flatten_table_to_lines(table, skip_header=False)
        elif "sku" in key and "mrp" in key:
            flatten = flatten_table_to_lines(table, skip_header=False)

        if flatten is None:
            continue

        _replace_element_with_paragraphs(doc, table._tbl, flatten)


def parse_swot_block(text: str) -> dict[str, list[str]] | None:
    headers = ("Strengths", "Weaknesses", "Opportunities", "Threats")
    start = text.find("6. SWOT Analysis")
    if start < 0:
        return None
    end = text.find("Suggested response", start)
    if end < 0:
        end = text.find("7. Business Model", start)
    if end < 0:
        return None

    buckets: dict[str, list[str]] = {h: [] for h in headers}
    current: str | None = None
    for raw in text[start:end].splitlines():
        line = raw.strip()
        if not line or line == "6. SWOT Analysis":
            continue
        matched = next((h for h in headers if line == h), None)
        if matched:
            current = matched
            continue
        if current:
            buckets[current].append(line)
    return buckets if any(buckets.values()) else None


def fix_swot_like_smallpdf(docx_path: str, pdf_path: str) -> None:
    import fitz

    pdf = fitz.open(pdf_path)
    swot_blocks: list[dict[str, list[str]]] = []
    for page in pdf:
        parsed = parse_swot_block(page.get_text())
        if parsed:
            swot_blocks.append(parsed)
    pdf.close()

    if not swot_blocks:
        return

    doc = Document(docx_path)
    body = doc.element.body
    block_idx = 0
    changed = False

    for child in list(body):
        if child.tag != qn("w:p"):
            continue

        merged = "".join(node.text or "" for node in child.iter(qn("w:t"))).strip()
        if "6. SWOT Analysis" not in merged:
            continue

        if block_idx >= len(swot_blocks):
            break

        buckets = swot_blocks[block_idx]
        block_idx += 1

        for node in child.iter(qn("w:t")):
            node.text = "6. SWOT Analysis"

        cursor = body.index(child) + 1
        while cursor < len(body):
            nxt = body[cursor]
            if nxt.tag == qn("w:p"):
                nxt_text = "".join(
                    node.text or "" for node in nxt.iter(qn("w:t"))
                ).strip()
                if nxt_text.startswith("Suggested response") or nxt_text.startswith(
                    "7. Business Model"
                ):
                    break
            body.remove(nxt)

        insert_at = cursor
        pairs = (
            ("Strengths", "Weaknesses", buckets["Strengths"], buckets["Weaknesses"]),
            (
                "Opportunities",
                "Threats",
                buckets["Opportunities"],
                buckets["Threats"],
            ),
        )

        for left_h, right_h, left_items, right_items in pairs:
            row_count = max(len(left_items), len(right_items), 1) + 1
            tbl = doc.add_table(rows=row_count, cols=2)
            tbl.style = "Table Grid"
            tbl.cell(0, 0).text = left_h
            tbl.cell(0, 1).text = right_h
            for run in tbl.cell(0, 0).paragraphs[0].runs:
                run.bold = True
            for run in tbl.cell(0, 1).paragraphs[0].runs:
                run.bold = True

            for i in range(max(len(left_items), len(right_items))):
                if i < len(left_items):
                    tbl.cell(i + 1, 0).text = left_items[i]
                if i < len(right_items):
                    tbl.cell(i + 1, 1).text = right_items[i]

            tbl_element = tbl._tbl
            tbl_element.getparent().remove(tbl_element)
            body.insert(insert_at, tbl_element)
            insert_at += 1

        changed = True

    if changed:
        doc.save(docx_path)


def _paragraph_in_ledger_table(paragraph: Paragraph) -> bool:
    node = paragraph._element.getparent()
    while node is not None:
        if node.tag == qn("w:tbl"):
            text = "".join(t.text or "" for t in node.iter(qn("w:t")))
            return any(marker in text for marker in LEDGER_MARKERS)
        node = node.getparent()
    return False


def _paragraph_in_table(paragraph: Paragraph) -> bool:
    node = paragraph._element.getparent()
    while node is not None:
        if node.tag == qn("w:tbl"):
            return True
        node = node.getparent()
    return False


def add_list_numbering(paragraph: Paragraph, num_id: int = 1, ilvl: int = 0) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    if p_pr.find(qn("w:numPr")) is not None:
        return

    p_style = p_pr.find(qn("w:pStyle"))
    if p_style is None:
        p_style = OxmlElement("w:pStyle")
        p_pr.insert(0, p_style)
    p_style.set(qn("w:val"), "ListParagraph")

    num_pr = OxmlElement("w:numPr")
    ilvl_el = OxmlElement("w:ilvl")
    ilvl_el.set(qn("w:val"), str(ilvl))
    num_id_el = OxmlElement("w:numId")
    num_id_el.set(qn("w:val"), str(num_id))
    num_pr.append(ilvl_el)
    num_pr.append(num_id_el)
    p_pr.append(num_pr)

    jc = p_pr.find(qn("w:jc"))
    if jc is None:
        jc = OxmlElement("w:jc")
        p_pr.append(jc)
    jc.set(qn("w:val"), "left")


NON_LIST_EXACT = {
    "Strategy Role Priority",
    "Segment Product Role Launch Phase",
    "SKU MRP (incl. GST) Role",
    "Brand Position Primary Model Primary Engine",
    "Strengths",
    "Weaknesses",
    "Opportunities",
    "Threats",
    "Sales funnel",
    "Recommended pricing",
    "Financial structure",
    "Base-case targets",
    "Core marketing message",
    "Recommended product architecture",
    "Formats",
}


def apply_native_lists(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        text = paragraph.text.strip()
        if not text:
            continue
        if HEADING_RE.match(text):
            continue
        if text in NON_LIST_EXACT:
            continue
        if _paragraph_in_table(paragraph):
            continue
        if _paragraph_in_ledger_table(paragraph):
            continue
        if text.startswith("Suggested response"):
            continue
        add_list_numbering(paragraph)


def dedupe_consecutive_paragraphs(doc: Document) -> None:
    previous = ""
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if text and text == previous:
            element = paragraph._element
            parent = element.getparent()
            if parent is not None:
                parent.remove(element)
            continue
        if text:
            previous = text
        else:
            previous = ""


def normalize_like_smallpdf(docx_path: str, pdf_path: str) -> None:
    fix_swot_like_smallpdf(docx_path, pdf_path)

    doc = Document(docx_path)
    split_merged_paragraphs(doc)
    rebuild_sections_from_pdf(doc, pdf_path)
    flatten_text_layout_tables(doc)
    split_kra_tables(doc, pdf_path)
    dedupe_consecutive_paragraphs(doc)
    apply_native_lists(doc)
    doc.save(docx_path)
