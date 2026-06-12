#!/usr/bin/env python3
"""PDF → DOCX using pdf2docx; optional Smallpdf transform for Founder Summary PDFs only."""
from __future__ import annotations

import io
import logging
import os
import re
import sys
import threading
import time
import zipfile
from pathlib import Path

SCRIPT_DIR = Path(__file__).resolve().parent
_PROGRESS_PCT = 0
_LAST_EMIT_MS = 0.0
MAX_IMAGE_PATH_PAGES = 150
LARGE_PDF_PAGES = 200
CHUNK_MIN_PAGES = 400
CHUNK_SIZE = 250
MAX_CHUNK_WORKERS = min(4, os.cpu_count() or 2)

_PROGRESS_SCOPE = {"base": 10, "span": 75}


def emit_progress(pct: int, *, force: bool = False) -> None:
    global _PROGRESS_PCT, _LAST_EMIT_MS
    pct = max(0, min(99, int(pct)))
    if pct <= _PROGRESS_PCT:
        return
    now = time.monotonic()
    if (
        not force
        and pct < 99
        and (now - _LAST_EMIT_MS) < 0.35
        and (pct - _PROGRESS_PCT) < 2
    ):
        return
    _PROGRESS_PCT = pct
    _LAST_EMIT_MS = now
    print(f"PROGRESS pct={pct}", file=sys.stderr, flush=True)


class Pdf2DocxProgressHandler(logging.Handler):
    """Map pdf2docx phase and page logs to monotonic PROGRESS pct lines."""

    _PHASE_PCT = {
        1: 12,  # Opening document
        2: 22,  # Analyzing document
        3: 32,  # Parsing pages
        4: 42,  # Creating pages
    }

    def emit(self, record: logging.LogRecord) -> None:
        msg = re.sub(r"\x1b\[[0-9;]*m", "", record.getMessage())

        phase = re.search(r"\[(\d+)/4\]\s+(\w+)", msg)
        if phase:
            emit_progress(self._PHASE_PCT.get(int(phase.group(1)), 15))
            return

        match = re.search(r"\((\d+)/(\d+)\)\s+Page", msg)
        if not match:
            return
        current = int(match.group(1))
        total = max(1, int(match.group(2)))
        base = _PROGRESS_SCOPE["base"]
        span = _PROGRESS_SCOPE["span"]
        emit_progress(base + int(span * current / total))


def install_progress_logging() -> None:
    handler = Pdf2DocxProgressHandler()
    handler.setLevel(logging.INFO)
    root = logging.getLogger()
    root.handlers = [handler]
    root.setLevel(logging.INFO)


def optimal_settings(page_count: int = 1, *, image_heavy: bool = False) -> dict:
    cpus = min(4, os.cpu_count() or 2)
    # pdf2docx pool workers on Windows re-import cv2 and hang or crash (spawn).
    use_mp = page_count > 4 and sys.platform != "win32"
    parse_tables = not image_heavy and page_count <= 200
    return {
        "multi_processing": use_mp,
        "cpu_count": cpus if use_mp else 1,
        "ignore_page_error": True,
        "delete_end_line_hyphen": True,
        "parse_lattice_table": parse_tables,
        "parse_stream_table": False,
        "list_not_table": False,
    }


def _progress_heartbeat(stop: threading.Event, ceiling: int = 93) -> None:
    """Slowly advance progress while pdf2docx runs without log updates."""
    while not stop.wait(2.0):
        if _PROGRESS_PCT >= ceiling:
            continue
        emit_progress(min(ceiling, _PROGRESS_PCT + 1), force=True)


def _sample_page_indices(page_count: int) -> list[int]:
    if page_count <= 60:
        return list(range(page_count))
    step = max(1, page_count // 25)
    return list(range(0, page_count, step))[:30]


def prepare_input_pdf(pdf_path: str) -> tuple[str, list[str]]:
    """Decrypt/repair PDF so pdf2docx can read it. Returns path + temp files to delete."""
    import fitz

    cleanup: list[str] = []
    password = os.environ.get("PDF_INPUT_PASSWORD", "")

    doc = fitz.open(pdf_path)
    try:
        if doc.is_encrypted:
            authenticated = False
            for pw in (password, ""):
                if pw is not None and doc.authenticate(pw):
                    authenticated = True
                    break
            if not authenticated:
                print(
                    "ERROR PASSWORD_REQUIRED This PDF is password-protected.",
                    file=sys.stderr,
                )
                return "", cleanup

        needs_clean = doc.is_encrypted
        if not needs_clean and len(doc) > 0:
            try:
                doc[0].get_text("words")
            except ValueError:
                needs_clean = True

        if needs_clean:
            work = os.path.join(os.path.dirname(pdf_path), "input-ready.pdf")
            doc.save(work, garbage=4, deflate=True)
            cleanup.append(work)
            return work, cleanup

        return pdf_path, cleanup
    finally:
        doc.close()


def docx_metrics(docx_path: str) -> dict[str, int]:
    import re

    with zipfile.ZipFile(docx_path) as archive:
        xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        media = len([n for n in archive.namelist() if n.startswith("word/media/")])

    chars = sum(len(m) for m in re.findall(r"<w:t[^>]*>([^<]*)</w:t>", xml))
    return {
        "tables": xml.count("<w:tbl>"),
        "drawings": xml.count("<w:drawing"),
        "media": media,
        "numPr": xml.count("<w:numPr"),
        "chars": chars,
    }


def validate_docx(docx_path: str, *, image_ok: bool = False) -> bool:
    """Fast validity check — no full-document char scan."""
    if not os.path.isfile(docx_path) or os.path.getsize(docx_path) < 1500:
        return False
    try:
        with zipfile.ZipFile(docx_path) as archive:
            xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
        has_text = "<w:t" in xml
        has_drawings = "<w:drawing" in xml
        has_tables = "<w:tbl>" in xml
        if has_text or has_drawings or has_tables:
            return True
        if image_ok and has_drawings and os.path.getsize(docx_path) >= 3000:
            return True
        return False
    except zipfile.BadZipFile:
        return False


def pdf_open_stats(pdf_path: str) -> tuple[int, int, bool, bool, bool]:
    """One pass: page_count, text_pages_sample, needs_transform, image_heavy, drawing_heavy."""
    import fitz

    doc = fitz.open(pdf_path)
    try:
        page_count = len(doc)
        if page_count == 0:
            return 0, 0, False, False, False

        text_pages = 0
        for i in range(min(4, page_count)):
            if len(doc[i].get_text("words")) > 12:
                text_pages += 1

        image_heavy = text_pages == 0

        heavy_draw_pages = 0
        quick_sample = min(4, page_count)
        for i in range(quick_sample):
            page = doc[i]
            drawings = len(page.get_drawings())
            words = len(page.get_text("words"))
            if drawings > 600 or (drawings > 250 and words < 25):
                heavy_draw_pages += 1

        drawing_heavy = heavy_draw_pages >= max(2, quick_sample - 1)
        if not drawing_heavy and page_count <= LARGE_PDF_PAGES:
            for i in _sample_page_indices(page_count):
                if i < quick_sample:
                    continue
                page = doc[i]
                drawings = len(page.get_drawings())
                words = len(page.get_text("words"))
                if drawings > 600 or (drawings > 250 and words < 25):
                    heavy_draw_pages += 1

            sampled = len(_sample_page_indices(page_count))
            drawing_heavy = heavy_draw_pages >= max(2, int(sampled * 0.35))

        needs_transform = False
        if page_count >= 15 and not drawing_heavy:
            p0 = doc[0].get_text()
            if "CONFIDENTIAL" in p0:
                if "Founder Summary" in p0 or "The786" in p0:
                    needs_transform = True
                else:
                    sample = "".join(
                        doc[i].get_text() for i in range(min(8, page_count))
                    )
                    needs_transform = (
                        "SWOT Analysis" in sample and "Business Model" in sample
                    )

        return page_count, text_pages, needs_transform, image_heavy, drawing_heavy
    finally:
        doc.close()


def convert_image_pdf_to_docx(pdf_path: str, docx_path: str) -> bool:
    """Fallback for image-only / scanned PDFs — one page image per Word page."""
    import fitz
    from docx import Document
    from docx.shared import Inches

    doc = fitz.open(pdf_path)
    try:
        if len(doc) == 0:
            return False
        word = Document()
        usable_width = Inches(6.5)
        total = len(doc)
        for page_index, page in enumerate(doc):
            emit_progress(12 + int(75 * (page_index + 1) / total))
            if page_index > 0:
                word.add_page_break()
            rect = page.rect
            zoom = min(2.0, 1600 / max(rect.width, 1))
            pix = page.get_pixmap(matrix=fitz.Matrix(zoom, zoom), alpha=False)
            png = pix.tobytes("png")
            word.add_picture(io.BytesIO(png), width=usable_width)
        word.save(docx_path)
    finally:
        doc.close()
    return validate_docx(docx_path, image_ok=True)


def merge_docx_files(paths: list[str], output_path: str) -> None:
    """Append DOCX bodies in order (text/layout PDFs)."""
    from copy import deepcopy

    from docx import Document

    if not paths:
        raise ValueError("No DOCX parts to merge")

    master = Document(paths[0])
    for part in paths[1:]:
        doc = Document(part)
        for child in list(doc.element.body):
            if child.tag.endswith("sectPr"):
                continue
            master.element.body.append(deepcopy(child))
    master.save(output_path)


def _convert_page_range(args: tuple[str, int, int, str, bool]) -> tuple[int, int, str]:
    """Worker: convert inclusive page range to a standalone DOCX."""
    pdf_path, start, end, out_path, image_heavy = args
    from pdf2docx import Converter

    pages = end - start + 1
    cv = Converter(pdf_path)
    try:
        cv.convert(
            out_path,
            start=start,
            end=end,
            **optimal_settings(pages, image_heavy=image_heavy),
        )
    finally:
        cv.close()
    return start, end, out_path


def _convert_page_range_subprocess(
    args: tuple[str, int, int, str, bool],
) -> tuple[int, int, str]:
    """Windows-safe chunk worker via isolated subprocess."""
    import subprocess

    pdf_path, start, end, out_path, _image_heavy = args
    range_script = SCRIPT_DIR / "pdf-to-docx-range.py"
    result = subprocess.run(
        [sys.executable, str(range_script), pdf_path, out_path, str(start), str(end)],
        capture_output=True,
        text=True,
    )
    if result.returncode != 0:
        err = (result.stderr or result.stdout or "").strip()
        raise RuntimeError(err or f"chunk {start}-{end} failed")
    return start, end, out_path


def _run_chunks(
    chunk_specs: list[tuple[str, int, int, str, bool]],
    page_count: int,
    *,
    parallel: bool,
) -> list[tuple[int, str]] | None:
    """Run chunk conversions; return sorted outputs or None on failure."""
    from concurrent.futures import ThreadPoolExecutor, as_completed

    total_chunks = len(chunk_specs)
    chunk_outputs: list[tuple[int, str]] = []
    completed_pages = 0
    worker = _convert_page_range_subprocess if parallel else _convert_page_range
    max_workers = min(MAX_CHUNK_WORKERS, total_chunks) if parallel else 1

    def on_done(start: int, end: int, out_path: str) -> bool:
        nonlocal completed_pages
        if not validate_docx(out_path):
            return False
        chunk_outputs.append((start, out_path))
        completed_pages += end - start + 1
        pct = 10 + int(75 * completed_pages / page_count)
        emit_progress(min(86, pct), force=True)
        return True

    if parallel and max_workers > 1:
        with ThreadPoolExecutor(max_workers=max_workers) as pool:
            futures = {pool.submit(worker, spec): spec for spec in chunk_specs}
            for future in as_completed(futures):
                spec = futures[future]
                emit_progress(10 + int(75 * spec[1] / page_count), force=True)
                try:
                    start, end, out_path = future.result()
                except Exception as exc:
                    print(f"WARN chunk failed: {exc}", file=sys.stderr)
                    return None
                if not on_done(start, end, out_path):
                    return None
    else:
        for idx, spec in enumerate(chunk_specs):
            start, end = spec[1], spec[2]
            emit_progress(10 + int(75 * start / page_count), force=True)
            print(
                f"CHUNK {idx + 1}/{total_chunks} pages {start + 1}-{end + 1}",
                file=sys.stderr,
                flush=True,
            )
            try:
                c_start, c_end, c_out = worker(spec)
            except Exception as exc:
                print(f"WARN chunk failed: {exc}", file=sys.stderr)
                return None
            if not on_done(c_start, c_end, c_out):
                return None

    chunk_outputs.sort(key=lambda item: item[0])
    return chunk_outputs


def convert_pdf_chunked(
    pdf_path: str,
    docx_path: str,
    page_count: int,
    *,
    image_heavy: bool,
    drawing_heavy: bool,
) -> bool:
    """Large PDFs: page-range conversion in chunks + DOCX merge."""
    import shutil
    import tempfile

    heavy = image_heavy or drawing_heavy
    chunk_size = 100 if page_count > 2000 else CHUNK_SIZE
    tmp = tempfile.mkdtemp(prefix="pdf2docx-chunk-")
    chunk_specs: list[tuple[str, int, int, str, bool]] = []
    try:
        for start in range(0, page_count, chunk_size):
            end = min(start + chunk_size - 1, page_count - 1)
            out_path = os.path.join(tmp, f"part_{start:05d}.docx")
            chunk_specs.append((pdf_path, start, end, out_path, heavy))

        total_chunks = len(chunk_specs)
        emit_progress(10, force=True)

        stop_heartbeat = threading.Event()
        heartbeat = threading.Thread(
            target=_progress_heartbeat, args=(stop_heartbeat, 86), daemon=True
        )
        heartbeat.start()
        try:
            use_parallel = total_chunks > 1
            chunk_outputs = _run_chunks(
                chunk_specs, page_count, parallel=use_parallel
            )
            if chunk_outputs is None:
                print("WARN retrying chunks sequentially", file=sys.stderr)
                chunk_outputs = _run_chunks(
                    chunk_specs, page_count, parallel=False
                )
            if chunk_outputs is None:
                return False
        finally:
            stop_heartbeat.set()
            heartbeat.join(timeout=1.0)

        emit_progress(88, force=True)
        merge_docx_files([path for _, path in chunk_outputs], docx_path)
        emit_progress(92, force=True)
        return validate_docx(docx_path)
    finally:
        shutil.rmtree(tmp, ignore_errors=True)


def convert_pdf_single(
    pdf_path: str,
    docx_path: str,
    page_count: int,
    *,
    image_heavy: bool,
    drawing_heavy: bool,
) -> None:
    from pdf2docx import Converter

    end_page = max(0, page_count - 1)
    _PROGRESS_SCOPE["base"] = 10
    _PROGRESS_SCOPE["span"] = 75

    cv = Converter(pdf_path)
    stop_heartbeat = threading.Event()
    heartbeat = threading.Thread(
        target=_progress_heartbeat, args=(stop_heartbeat,), daemon=True
    )
    try:
        heartbeat.start()
        cv.convert(
            docx_path,
            start=0,
            end=end_page,
            **optimal_settings(page_count, image_heavy=image_heavy or drawing_heavy),
        )
    finally:
        stop_heartbeat.set()
        heartbeat.join(timeout=1.0)
        cv.close()


def convert_pdf(pdf_path: str, docx_path: str) -> bool:
    work_pdf, temp_files = prepare_input_pdf(pdf_path)
    if not work_pdf:
        return False

    try:
        return _convert_pdf_inner(work_pdf, docx_path)
    finally:
        for path in temp_files:
            try:
                os.remove(path)
            except OSError:
                pass


def _convert_pdf_inner(pdf_path: str, docx_path: str) -> bool:
    emit_progress(3)
    page_count, _text_pages, needs_transform, image_heavy, drawing_heavy = (
        pdf_open_stats(pdf_path)
    )
    emit_progress(8)

    if page_count == 0:
        return False

    use_image_path = (image_heavy or drawing_heavy) and page_count <= MAX_IMAGE_PATH_PAGES
    if use_image_path:
        emit_progress(12)
        ok = convert_image_pdf_to_docx(pdf_path, docx_path)
        emit_progress(95)
        return ok

    if page_count >= CHUNK_MIN_PAGES:
        ok = convert_pdf_chunked(
            pdf_path,
            docx_path,
            page_count,
            image_heavy=image_heavy,
            drawing_heavy=drawing_heavy,
        )
        if not ok:
            return False
        emit_progress(88)
    else:
        convert_pdf_single(
            pdf_path,
            docx_path,
            page_count,
            image_heavy=image_heavy,
            drawing_heavy=drawing_heavy,
        )
        emit_progress(88)

    if needs_transform:
        sys.path.insert(0, str(SCRIPT_DIR))
        from smallpdf_transform import transform_to_smallpdf

        transform_to_smallpdf(docx_path, pdf_path)
        emit_progress(95)
        return validate_docx(docx_path, image_ok=False)

    if validate_docx(docx_path, image_ok=image_heavy):
        return True

    if image_heavy:
        return convert_image_pdf_to_docx(pdf_path, docx_path)

    return False


def main() -> int:
    if len(sys.argv) != 3:
        print("Usage: pdf-to-docx.py <input.pdf> <output.docx>", file=sys.stderr)
        return 1

    import logging

    install_progress_logging()

    pdf_path, docx_path = sys.argv[1], sys.argv[2]

    try:
        from pdf2docx import Converter  # noqa: F401
    except ImportError:
        print("pdf2docx not installed. Run: pip install pdf2docx", file=sys.stderr)
        return 2

    if not convert_pdf(pdf_path, docx_path):
        print("Conversion produced invalid DOCX", file=sys.stderr)
        return 3

    size_kb = os.path.getsize(docx_path) // 1024
    emit_progress(99)
    print(f"OK size_kb={size_kb}", file=sys.stderr, flush=True)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
