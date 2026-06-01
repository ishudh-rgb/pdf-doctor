#!/usr/bin/env python3
from __future__ import annotations

import sys
from docx import Document
from docx.oxml.ns import qn


def has_num(p) -> bool:
    ppr = p._p.find(qn("w:pPr"))
    return ppr is not None and ppr.find(qn("w:numPr")) is not None


def section(path: str, start: str, end: str | None, limit: int = 40) -> None:
    doc = Document(path)
    active = False
    rows: list[tuple[bool, str]] = []
    for p in doc.paragraphs:
        t = p.text.strip()
        if start in t:
            active = True
        if not active:
            continue
        if end and t.startswith(end) and start not in t:
            break
        if t:
            rows.append((has_num(p), t))
        if len(rows) >= limit:
            break

    print(f"\n=== {path.split(chr(92))[-1]} : {start} ===")
    for num, t in rows:
        mark = "B" if num else " "
        print(f"{mark} | {t[:90]}")


def totals(path: str) -> None:
    doc = Document(path)
    n = b = empty_b = 0
    for p in doc.paragraphs:
        t = p.text.strip()
        if not t:
            continue
        n += 1
        if has_num(p):
            b += 1
            if len(t) < 3:
                empty_b += 1
    print(f"{path.split(chr(92))[-1]}: paras={n} bullets={b} empty_bullet_text={empty_b}")


if __name__ == "__main__":
    ref = sys.argv[1]
    ours = sys.argv[2]
    for path in (ref, ours):
        totals(path)
    for start, end in [
        ("Price Market Gap", "SWOT"),
        ("Problem Statement", "Solution"),
        ("Suggested response", "Business Model"),
        ("Gap ka simple version", "SWOT"),
    ]:
        section(ref, start, end)
        section(ours, start, end)
