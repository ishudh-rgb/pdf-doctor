#!/usr/bin/env python3
"""
Build DOCX matching Smallpdf result.docx layout for merged (3).pdf.

Smallpdf uses:
- Headings WITHOUT numeric prefix (Problem Statement, not 1. Problem Statement)
- Tables for brand, products, marketing, pricing, KRA
- Single 2-column SWOT table
- Separate paragraphs for formats / market-gap bullets / sales funnel
"""
from __future__ import annotations

import re
from typing import Iterable

import fitz
from docx import Document
from docx.shared import Pt
from docx.table import Table


def pdf_text(pdf_path: str) -> str:
    doc = fitz.open(pdf_path)
    text = "".join(page.get_text() for page in doc)
    doc.close()
    return text


def clean_lines(block: str) -> list[str]:
    lines: list[str] = []
    for raw in block.splitlines():
        line = raw.strip()
        if not line or line == "CONFIDENTIAL":
            continue
        lines.append(line)
    return lines


def section(full: str, start: str, end: str) -> list[str]:
    s = full.find(start)
    if s < 0:
        return []
    e = full.find(end, s + len(start))
    if e < 0:
        e = len(full)
    return clean_lines(full[s:e])


def strip_heading_number(line: str) -> str:
    return re.sub(r"^\d+\.\s+", "", line).strip()


def add_heading(doc: Document, text: str) -> None:
    doc.add_paragraph(strip_heading_number(text))


def add_bullets(doc: Document, lines: Iterable[str]) -> None:
    for line in lines:
        doc.add_paragraph(line, style="List Paragraph")


def add_table_grid(doc: Document, rows: list[list[str]]) -> Table:
    if not rows:
        return doc.add_table(1, 1)
    cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    for ri, row in enumerate(rows):
        for ci in range(cols):
            table.cell(ri, ci).text = row[ci] if ci < len(row) else ""
    return table


def parse_product_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for line in lines:
        if line.startswith("Phase "):
            parts = line.split(None, 3)
            if len(parts) >= 4:
                rows.append([parts[0], parts[1], parts[2], parts[3]])
            continue
        m = re.match(
            r"^(Men|Women|Unisex|Sampler)\s+(.+?)\s+(Phase \d+)$",
            line,
        )
        if m:
            segment, rest, phase = m.groups()
            role = rest
            product = rest
            if " Phase " in rest:
                product, _, phase2 = rest.rpartition(" Phase ")
                phase = "Phase " + phase2.strip()
            rows.append([segment, product, role, phase])
            continue
        # PDF line format: Men Blue Vetiver Musk Fresh / office-safe hero Phase 1
        m2 = re.match(r"^(Men|Women|Unisex|Sampler)\s+(.+)\s+(Phase \d+)$", line)
        if m2:
            segment, middle, phase = m2.groups()
            rows.append([segment, middle, middle, phase])
    return rows


def parse_marketing_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    priorities = ("High", "Very High")
    i = 0
    while i < len(lines):
        line = lines[i]
        if line in {"Strategy", "Role", "Priority"}:
            i += 1
            continue
        if line.startswith("Core marketing message"):
            break
        strategy = line
        role = lines[i + 1] if i + 1 < len(lines) else ""
        priority = lines[i + 2] if i + 2 < len(lines) else ""
        if priority not in priorities and role in priorities:
            priority, role = role, ""
        if role in priorities:
            rows.append([strategy, "", role])
            i += 2
        elif priority in priorities:
            rows.append([strategy, role, priority])
            i += 3
        else:
            i += 1
    return rows


def parse_pricing_rows(lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    i = 0
    while i < len(lines):
        line = lines[i]
        if line in {"SKU", "MRP (incl. GST)", "Role", "Recommended pricing"}:
            i += 1
            continue
        if line == "Financial structure":
            break
        if i + 2 < len(lines) and lines[i + 1].startswith("₹"):
            rows.append([line, lines[i + 1], lines[i + 2]])
            i += 3
        else:
            i += 1
    return rows


def parse_kra_rows(full: str) -> list[list[str]]:
    block = section(full, "Recommended role:", "Suggested KPI scorecard weight")
    lines = [ln for ln in block if ln not in {"KRA", "What they own", "Core KPIs"}]
    if lines and lines[0].startswith("Recommended role:"):
        lines = lines[1:]

    rows: list[list[str]] = []
    buf: list[str] = []
    for line in lines:
        buf.append(line)
        if len(buf) == 3:
            rows.append(buf)
            buf = []
    if buf:
        while len(buf) < 3:
            buf.append("")
        rows.append(buf)
    return rows


def parse_swot(full: str) -> tuple[list[str], list[str]]:
    block = section(full, "6. SWOT Analysis", "Suggested response")
    left: list[str] = []
    right: list[str] = []
    side = "left"
    for line in block:
        if line == "6. SWOT Analysis":
            continue
        if line == "Strengths":
            side = "left"
            continue
        if line == "Weaknesses":
            side = "right"
            continue
        if line == "Opportunities":
            side = "left"
            continue
        if line == "Threats":
            side = "right"
            continue
        if side == "left":
            left.append(line)
        else:
            right.append(line)
    return left, right


def build_copy(doc: Document, full: str) -> None:
    # Title
    doc.add_paragraph("The786 — Founder Summary Sheet")
    doc.add_paragraph(
        "Problem → Solution → Products → SWOT → Business Model → Marketing → "
        "Sales → Financials → GTM → Co-founder KRA/KPI"
    )

    # Brand table 1x3
    brand = section(full, "Brand Position", "1. Problem Statement")
    if len(brand) >= 3:
        add_table_grid(
            doc,
            [
                [
                    "Brand Position",
                    "Primary Model",
                    "Primary Engine",
                ],
                brand[:3],
            ],
        )

    # Problem
    add_heading(doc, "1. Problem Statement")
    problem = section(full, "1. Problem Statement", "2. Solution")
    if problem:
        add_bullets(doc, problem[1:])

    # Solution
    add_heading(doc, "2. Solution")
    solution = section(full, "2. Solution", "3. Product")
    if solution:
        for line in solution[1:]:
            doc.add_paragraph(line)

    # Products
    add_heading(doc, "3. Products")
    products = section(full, "Recommended product architecture", "Formats")
    doc.add_paragraph("Recommended product architecture")
    prod_lines = [ln for ln in products if ln != "Recommended product architecture"]
    prod_rows = parse_product_rows(prod_lines)
    if prod_rows:
        add_table_grid(
            doc,
            [["Segment", "Product", "Role", "Launch Phase"], *prod_rows],
        )

    formats = section(full, "Formats", "4. Market Gap")
    if formats:
        doc.add_paragraph("Formats")
        for line in formats[1:]:
            doc.add_paragraph(line)

    # Market gap
    add_heading(doc, "4. Market Gap")
    gap = section(full, "4. Market Gap", "5. Price Market Gap")
    if gap:
        for line in gap[1:]:
            doc.add_paragraph(line)

    # Price
    add_heading(doc, "5. Price Market Gap")
    price = section(full, "5. Price Market Gap", "6. SWOT Analysis")
    if price:
        for line in price[1:]:
            doc.add_paragraph(line)

    # SWOT
    add_heading(doc, "6. SWOT Analysis")
    s_left, s_right = parse_swot(full)
    w_left, w_right = [], []
    block = section(full, "6. SWOT Analysis", "Suggested response")
    mode = None
    col_a: list[str] = []
    col_b: list[str] = []
    for line in block:
        if line in {"6. SWOT Analysis", "Strengths", "Opportunities"}:
            mode = "a"
            continue
        if line in {"Weaknesses", "Threats"}:
            mode = "b"
            continue
        if mode == "a":
            col_a.append(line)
        elif mode == "b":
            col_b.append(line)

    max_rows = max(len(col_a), len(col_b), 1)
    swot_rows = [["Strengths", "Weaknesses"]]
    for i in range(max_rows):
        swot_rows.append(
            [
                col_a[i] if i < len(col_a) else "",
                col_b[i] if i < len(col_b) else "",
            ]
        )
    add_table_grid(doc, swot_rows)

    # Business model block
    biz = section(full, "Suggested response", "8. Marketing Model")
    if biz:
        for line in biz:
            doc.add_paragraph(line)

    # Marketing
    add_heading(doc, "8. Marketing Model")
    mkt = section(full, "8. Marketing Model", "9. Sales Model")
    mkt_rows = parse_marketing_rows(mkt[1:] if mkt else [])
    if mkt_rows:
        add_table_grid(doc, [["Strategy", "Role", "Priority"], *mkt_rows])
    tail = section(full, "Core marketing message", "9. Sales Model")
    for line in tail:
        doc.add_paragraph(line)

    # Sales (no numbered heading in Smallpdf)
    sales = section(full, "9. Sales Model", "10. Financial Modeling")
    if sales:
        for line in sales[1:]:
            doc.add_paragraph(line)

    # Financial
    add_heading(doc, "10. Financial Modeling")
    fin = section(full, "Recommended pricing", "11. GTM")
    if fin:
        doc.add_paragraph("Recommended pricing")
        price_rows = parse_pricing_rows(fin)
        if price_rows:
            add_table_grid(
                doc,
                [["SKU", "MRP (incl. GST)", "Role"], *price_rows],
            )
        for line in fin:
            if line == "Recommended pricing" or line.startswith("Discovery Set") and "₹" in line:
                continue
            if line in {"SKU", "MRP (incl. GST)", "Role"}:
                continue
            if any(line.startswith(x) for x in ("Discovery Set", "10ml Travel", "50ml Hero", "B2B Gifting")) and "₹" in line:
                continue
            if line == "Financial structure" or line.startswith("Model revenue") or "=" in line or line.startswith("Bootstrap") or line.startswith("Better-prepared") or line.startswith("Annual") or line.startswith("Track:"):
                doc.add_paragraph(line)


def build_docx(pdf_path: str, docx_path: str) -> None:
    full = pdf_text(pdf_path)
    doc = Document()

    # Build from each copy in PDF (split on Problem Statement after first)
    copies = re.split(r"(?=1\. Problem Statement)", full)
    copies = [c for c in copies if "1. Problem Statement" in c]

    if not copies:
        copies = [full]

    for idx, copy_text in enumerate(copies):
        if idx > 0:
            doc.add_page_break()
        build_copy(doc, copy_text)

    doc.save(docx_path)


if __name__ == "__main__":
    import sys

    if len(sys.argv) != 3:
        print("Usage: build_smallpdf_docx.py input.pdf output.docx")
        raise SystemExit(1)
    build_docx(sys.argv[1], sys.argv[2])
    print("OK")
